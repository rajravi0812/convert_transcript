import re
import streamlit as st
from docx import Document
from docx.shared import RGBColor
from io import BytesIO


# -----------------------------
# Core processing functions
# -----------------------------
def clean_srt_and_structure(srt_text: str):
    """Clean transcript/SRT, detect headings, lists, and paragraphs."""

    raw_lines = srt_text.splitlines()
    ts_pattern = re.compile(r'^\d\d:\d\d:\d\d,\d+ --> \d\d:\d\d:\d\d,\d+')
    index_line = re.compile(r'^\s*\d+\s*$')

    cleaned_lines = []
    for line in raw_lines:
        line = line.strip()
        if not line or index_line.match(line) or ts_pattern.match(line):
            continue

        # --- Headings ---
        if re.match(r"^\d+(\.\d+)*\s+.+", line):  # e.g. "1.1.1 Something"
            cleaned_lines.append(("main_heading", line))
        elif line.endswith(":"):  # e.g. "Max Tokens:"
            if any(char.isdigit() for char in line):
                cleaned_lines.append(("main_heading", line))
            else:
                cleaned_lines.append(("sub_heading", line))
        elif re.search(r"\bModule\s*\d+", line, re.IGNORECASE):  # Module-based heading
            # Capture heading up to first "."
            m = re.match(r'^(.*?\.)\s*(.*)$', line)
            if m:
                cleaned_lines.append(("main_heading", m.group(1).strip()))
                if m.group(2).strip():
                    cleaned_lines.append(("text", m.group(2).strip()))
            else:
                cleaned_lines.append(("main_heading", line.strip()))

        # --- Bullets ---
        elif line.startswith(("-", "*", "·")):
            if line.endswith(":"):  # bullet ending with ":" → treat as sub heading
                cleaned_lines.append(("sub_heading", line.lstrip("-*· ").strip()))
            else:
                cleaned_lines.append(("list", line.lstrip("-*· ").strip()))

        # --- Normal text ---
        else:
            cleaned_lines.append(("text", line))

    # Merge lines into sentences
    merged = []
    buffer = ""
    for ttype, content in cleaned_lines:
        if ttype != "text":  # flush before heading/list
            if buffer:
                merged.append(("text", buffer.strip()))
                buffer = ""
            merged.append((ttype, content))
        else:
            buffer += " " + content
            if re.search(r"[.!?]$", content):  # end of sentence
                merged.append(("text", buffer.strip()))
                buffer = ""
    if buffer:
        merged.append(("text", buffer.strip()))

    return merged


def export_to_word(structured, output_file="output.docx") -> BytesIO:
    """Export structured text to Word with:
       - main headings (bold + underline + black)
       - sub headings (bold + black, no underline)
       - bullet points (only part before ':' underlined if present)
       - normal paragraphs
    """

    doc = Document()
    for ttype, content in structured:
        if ttype == "main_heading":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.underline = True
            run.font.color.rgb = RGBColor(0, 0, 0)

        elif ttype == "sub_heading":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.bold = True
            run.font.underline = False
            run.font.color.rgb = RGBColor(0, 0, 0)

        elif ttype == "list":
            p = doc.add_paragraph(style="ListBullet")

            if ":" in content:
                before, after = content.split(":", 1)
                run1 = p.add_run(before + ":")
                run1.font.underline = True
                run2 = p.add_run(after)  # text after colon (no underline)
            else:
                run = p.add_run(content)
                run.font.underline = True

        else:  # normal paragraph
            doc.add_paragraph(content)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def structured_to_text(structured):
    """Convert structured output back to plain text with markers."""
    out_lines = []
    for ttype, content in structured:
        if ttype == "main_heading":
            out_lines.append(f"[MAIN HEADING] {content}")
        elif ttype == "sub_heading":
            out_lines.append(f"[SUB HEADING] {content}")
        elif ttype == "list":
            out_lines.append(f"- {content}")
        else:
            out_lines.append(content)
    return "\n".join(out_lines)


# -----------------------------
# Streamlit UI
# -----------------------------
st.set_page_config(page_title="SRT/Transcript Cleaner → TXT/Word", layout="centered")
st.title("📖 SRT/Transcript Cleaner → TXT / Word")

st.write(
    "Paste your transcript or SRT text below. This tool will remove timestamps and numeric counters, "
    "detect main headings (with digits), sub headings (without digits or ending with ':'), detect bullet lists, "
    "and put each sentence into its own paragraph."
)

sample = st.toggle("Load sample")
default_text = ""
if sample:
    default_text = """1
00:00:00,600 --> 00:00:05,569
1.1.1 Definition, Importance of Prompt Engineering in Software Development:
At its core, prompt engineering is the practice of designing input queries for AI...
2
00:00:05,819 --> 00:00:11,496
Max Tokens:
· Test Methods: Run automated tests
· Controls length: prevents overflow
· Prevents overflow
3
00:00:12,000 --> 00:00:17,000
Welcome to Module 1: Introduction to Prompt Engineering for Developers.
"""

raw = st.text_area("Paste SRT or transcript text", value=default_text, height=300, placeholder="Paste here...")

if st.button("Convert"):
    structured = clean_srt_and_structure(raw)
    output_txt = structured_to_text(structured)

    st.subheader("📝 Result (TXT)")
    st.text_area("Cleaned & Organized Text", value=output_txt, height=300)

    st.download_button(
        label="⬇️ Download .txt",
        data=output_txt.encode("utf-8"),
        file_name="cleaned_transcript.txt",
        mime="text/plain"
    )

    word_buffer = export_to_word(structured)
    st.download_button(
        label="⬇️ Download .docx",
        data=word_buffer,
        file_name="cleaned_transcript.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.markdown("---")
st.caption("✅ Removes timestamps & counters\n"
           "✅ Main Headings (digit-based) = bold + underline + black\n"
           "✅ Sub Headings (no digits OR ending with ':') = bold + black\n"
           "✅ Bullet points = only text before ':' underlined\n"
           "✅ Bullet ending with ':' → treated as Sub Heading\n"
           "✅ Each sentence becomes its own paragraph\n"
           "✅ Export available as TXT or Word")
