import re
import streamlit as st
from docx import Document
from io import BytesIO

# -----------------------------
# Core processing functions
# -----------------------------
def split_into_sentences(text: str):
    """Split text into sentences while keeping punctuation. Avoid splitting on ellipses (...)."""
    text = re.sub(r'\s+', ' ', text).strip()
    if not text:
        return []
    text = text.replace("...", " …ELLIPSIS… ")
    parts = re.split(r'(?<=[.!?])\s+', text)
    sentences = []
    for p in parts:
        p = p.replace("…ELLIPSIS…", "...").strip()
        if p:
            sentences.append(p)
    return sentences

def flush_paragraphs(sentences, sentences_per_paragraph: int = 1):
    """Each sentence becomes its own paragraph."""
    paragraphs = []
    for sentence in sentences:
        if sentence.strip():
            paragraphs.append(sentence.strip())
    return paragraphs

def clean_srt_and_structure(
    srt_text: str,
    detect_headings: bool = True,
    markdown_headings: bool = False,
):
    """Main cleaning logic for transcript/SRT"""
    raw_lines = srt_text.splitlines()

    ts_pattern = re.compile(r'\b\d{2}:\d{2}:\d{2},\d{3}\s*-->\s*\d{2}:\d{2}:\d{2},\d{3}\b')
    index_line = re.compile(r'^\s*\d+\s*$')

    merged_text = []
    for line in raw_lines:
        if index_line.match(line):
            continue
        if ts_pattern.search(line):
            continue
        cleaned = re.sub(r'\s+', ' ', line).strip()
        if cleaned:
            merged_text.append(cleaned)

    joined_text = " ".join(merged_text)
    all_sentences = split_into_sentences(joined_text)

    out_blocks = []
    current_heading = ""
    sentence_buffer = []

    def push_section():
        nonlocal sentence_buffer, current_heading, out_blocks
        if not sentence_buffer:
            return
        paragraphs = flush_paragraphs(sentence_buffer, 1)  # always 1 sentence per paragraph
        if current_heading:
            if markdown_headings:
                out_blocks.append(f"## {current_heading.strip()}")
            else:
                out_blocks.append(current_heading.strip().upper())
                out_blocks.append("=" * len(current_heading.strip()))
        for p in paragraphs:
            out_blocks.append(p)
            out_blocks.append("")  # blank line
        sentence_buffer = []

    for sentence in all_sentences:
        if detect_headings and ":" in sentence:
            head, tail = sentence.split(":", 1)
            if len(head.strip()) >= 8 or len(head.strip().split()) >= 2:
                push_section()
                current_heading = head.strip()
                remainder = tail.strip()
                if remainder:
                    sentence_buffer.extend(split_into_sentences(remainder))
                continue
        sentence_buffer.append(sentence)

    push_section()
    final_text = "\n".join(out_blocks).strip() + "\n"
    return final_text

def export_to_word(text: str) -> BytesIO:
    """Export formatted text to Word with bold headings."""
    doc = Document()
    lines = text.splitlines()

    for i, line in enumerate(lines):
        if not line.strip():
            continue
        if line.startswith("## "):
            doc.add_heading(line.replace("##", "").strip(), level=2)
        elif i + 1 < len(lines) and set(lines[i + 1]) == {"="}:
            doc.add_heading(line.strip(), level=1)
        else:
            doc.add_paragraph(line)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# -----------------------------
# Streamlit UI
# -----------------------------
st.set_page_config(page_title="SRT/Transcript Cleaner → TXT/Word", layout="centered")
st.title("📖 SRT/Transcript Cleaner → TXT / Word")

st.write(
    "Paste your transcript or SRT text below. This tool will remove timestamps and numeric counters, "
    "turn `Heading:` patterns into headings, and put each sentence into its own paragraph."
)

with st.sidebar:
    st.header("⚙️ Options")
    detect_headings = st.checkbox("Treat `text:` as a heading", value=True)
    markdown_headings = st.checkbox("Use Markdown headings (##)", value=False)

sample = st.toggle("Load sample")
default_text = ""
if sample:
    default_text = """1
00:00:00,600 --> 00:00:05,569
Welcome to Module 7: In this module we will cover X, Y, Z.
2
00:00:05,819 --> 00:00:11,496
This is a sample line. Here is another sentence. And one more. Final sentence!
"""

raw = st.text_area("Paste SRT or transcript text", value=default_text, height=300, placeholder="Paste here...")

if st.button("Convert"):
    output = clean_srt_and_structure(
        raw,
        detect_headings=detect_headings,
        markdown_headings=markdown_headings,
    )

    st.subheader("📝 Result (TXT)")
    st.text_area("Cleaned & Organized Text", value=output, height=300)

    st.download_button(
        label="⬇️ Download .txt",
        data=output.encode("utf-8"),
        file_name="cleaned_transcript.txt",
        mime="text/plain"
    )

    word_buffer = export_to_word(output)
    st.download_button(
        label="⬇️ Download .docx",
        data=word_buffer,
        file_name="cleaned_transcript.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.markdown("---")
st.caption("✅ Each sentence is placed in its own paragraph.\n"
           "✅ Headings are auto-highlighted.\n"
           "✅ Export available as TXT or Word (with styled headings).")
